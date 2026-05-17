import os
import re
from PyPDF2 import PdfReader
from docx import Document

class ResumeParser:
    """Parse resume files (PDF and DOCX) and extract text"""
    
    @staticmethod
    def parse_file(file_path):
        """Parse resume file and return extracted text"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return ResumeParser._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return ResumeParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    @staticmethod
    def _parse_pdf(file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        return text
    
    @staticmethod
    def _parse_docx(file_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
        return text
    
    @staticmethod
    def extract_contact_info(text):
        """Extract email, phone, and name from resume text"""
        contact_info = {
            'email': None,
            'phone': None,
            'name': None
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group(0)
        
        # Extract phone
        phone_pattern = r'\b(?:\+?1[-.]?)?\(?([0-9]{3})\)?[-.]?([0-9]{3})[-.]?([0-9]{4})\b'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group(0)
        
        # Extract name (first line, assuming it's the name)
        lines = text.split('\n')
        if lines:
            first_line = lines[0].strip()
            if len(first_line) < 100 and len(first_line.split()) <= 4:
                contact_info['name'] = first_line
        
        return contact_info